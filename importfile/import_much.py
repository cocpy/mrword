import os
import re
import datetime
import logging


def add_word(word):
    """添加有效单词"""
    logging.info('import youdao')
    from utils.youdao import ReptileYouDao
    from db.sqlite3_connect import insert_data
    logging.info('import finish')
    if len(word) > 2:
        # s = s.lower().replace(" ", "")
        logging.info('word:' + word)
        context = ReptileYouDao.get_base_message(word)
        error_code = context['error_code']
        logging.info(error_code)
        if error_code == 0:
            now = datetime.datetime.now().strftime("%Y-%m-%d %H:%M:%S")
            explains = context['explains']
            phonetic = context['phonetic']
            examples_en = context['examples_en']
            examples_cn = context['examples_cn']
            # 音标中可能含有'
            # sql = "INSERT INTO WORD (word,explains,phonetic,examples_en,examples_cn,create_time) " \
            #       "VALUES (\'" + word + "\',\'" + explains + "\',\'" + phonetic + "\',\'" + examples_en + "\',\'" + examples_cn + "\',\'" + now + "\')"
            sql = 'INSERT INTO WORD (word,explains,phonetic,examples_en,examples_cn,create_time) ' \
                  'VALUES (\"' + word + '\",\"' + explains + '\",\"' + phonetic + '\",\"' + examples_en + '\",\"' + examples_cn + '\",\"' + now + '\")'
            insert_data(sql)


def sub_cn(content):
    """通过正则替换 中文字符 为 ,"""

    # 替换中文
    pattern_cn = re.compile(r'[\u4e00-\u9fa5]')
    english = re.sub(pattern_cn, ',', content)
    # 替换数字
    pattern_num = re.compile(r'[0-9]')
    english = re.sub(pattern_num, ',', english)
    # 替换特殊字母
    r = ['n.', 'a.', 'adj.', 'adv.', 'vi.', 'vt.', 'prep.', 'ad.']
    for r_r in r:
        english = english.replace(r_r, '')
    # 替换符号
    r = ']“”。，‘’：；——{}【】、|《》？·~！@#￥$%……&*（）/\\><()^_!,-.;:?+="\'['
    english = re.sub(r'[{}]+'.format(r), ',', english)
    # 替换换行
    english = english.replace("\r\n", ",")
    # 替换空格
    english = english.lower().replace(" ", "").strip()
    return english


def common_handle(content):
    """各种文件共同的处理方法，传入str字符串"""
    content = str(content)
    logging.info('content:' + content)
    # 去除多余字符
    file_content = sub_cn(content)
    # 按','拆分
    str_list = file_content.split(',')
    logging.info('we are ready to add_word()')
    for s in str_list:
        # 添加有效单词
        add_word(s)


def pdf_read(file):
    """读取pdf文件"""

    # pip install pdfminer3k
    from pdfminer.pdfparser import PDFParser, PDFDocument
    from pdfminer.pdfinterp import PDFResourceManager, PDFPageInterpreter
    from pdfminer.converter import PDFPageAggregator
    from pdfminer.layout import LTTextBoxHorizontal, LAParams
    from pdfminer.pdfinterp import PDFTextExtractionNotAllowed

    fp = open(file, 'rb')
    # 用文件对象创建一个PDF文档分析器
    parser = PDFParser(fp)
    # 创建一个PDF文档
    doc = PDFDocument()
    # 连接分析器，与文档对象
    parser.set_document(doc)
    doc.set_parser(parser)

    # 提供初始化密码，如果没有密码，就创建一个空的字符串
    doc.initialize()
    # 检测文档是否提供txt转换，不提供就忽略
    if not doc.is_extractable:
        raise PDFTextExtractionNotAllowed
    else:
        # 创建PDF，资源管理器，来共享资源
        rsrcmgr = PDFResourceManager()
        # 创建一个PDF设备对象
        laparams = LAParams()
        device = PDFPageAggregator(rsrcmgr, laparams=laparams)
        # 创建一个PDF解释其对象
        interpreter = PDFPageInterpreter(rsrcmgr, device)
        # 循环遍历列表，每次处理一个page内容
        # doc.get_pages() 获取page列表
        for page in doc.get_pages():
            interpreter.process_page(page)
            # 接受该页面的LTPage对象
            layout = device.get_result()
            # 这里layout是一个LTPage对象 里面存放着 这个page解析出的各种对象
            # 一般包括LTTextBox, LTFigure, LTImage, LTTextBoxHorizontal 等等
            # 想要获取文本就获得对象的text属性，
            for x in layout:
                if isinstance(x, LTTextBoxHorizontal):
                    results = x.get_text()
                    common_handle(results)


def excel_read(file):
    """读取xlsx文件"""
    # pip install xlrd
    import xlrd

    workbook = xlrd.open_workbook(file)
    for sheet in workbook.sheets():
        row = sheet.nrows  # sheet行数
        col = sheet.ncols  # sheet列数
        for i in range(row):
            for j in range(col):
                value = str(sheet.cell_value(i, j))
                # 调用公共的处理方法
                common_handle(value)


def txt_read(file):
    """读取txt文件"""
    logging.info('txt_read will run')
    import chardet

    logging.info('txt_read is running')
    with open(file, "rb") as f:
        logging.info('file has opened')
        file_data = f.read()
        result = chardet.detect(file_data)
        # 解码
        file_content = file_data.decode(encoding=result['encoding'])
        # 调用公共的处理方法
        logging.info('is ready to run common_handle')
        common_handle(file_content)


def docx_read(file):
    """读取docx文件"""
    logging.info('docx_read is running')
    # pip install python_docx
    from docx import Document
    logging.info('Document has been imported')
    doc = Document(file)
    logging.info('Document finish')
    # 读取每一行文本，一个换行符算一行
    for paragraph in doc.paragraphs:
        # 调用公共的处理方法
        common_handle(paragraph.text)

    # 读取表格的方法，按需添加
    # for table in doc.tables:
    #     for row in table.rows:
    #         for cell in row.cells:
    #             print(cell.text)


def doc_read(path):
    """将 .doc 文件转成 .docx，再读取docx"""
    return True
    # 仅可在Windows下使用，暂时注释
    # import win32com.client
    #
    # w = win32com.client.Dispatch('Word.Application')
    # w.Visible = 0
    # w.DisplayAlerts = 0
    #
    # doc = w.Documents.Open(path)
    # newpath = os.path.splitext(path)[0] + '.docx'
    # doc.SaveAs(newpath, 12, False, "", True, "", False, False, False, False)
    # doc.Close()
    # w.Quit()
    # os.remove(path)
    # docx_read(newpath)


def import_file(file):
    """导入文件"""

    logging.info('start import file!!!')
    # 获取文件后缀
    postfix = os.path.splitext(file)[-1][1:]
    logging.info('postfix:'+postfix)
    if postfix == 'txt':
        logging.info('It\'s ready to read txt')
        txt_read(file)
    # elif postfix == 'doc':
    #     doc_read(file)
    elif postfix == 'docx':
        logging.info('It\'s ready to read docx')
        docx_read(file)
    elif postfix == 'xlsx':
        excel_read(file)
    elif postfix == 'pdf':
        pdf_read(file)
    else:
        logging.info('Can\'t import %s file' % postfix)


if __name__ == '__main__':
    # docx_read()
    # txt_read()
    # excel_read()
    # sub_cn()
    # pdf_read()
    # check_word('hello')
    import_file('test.tqt')
